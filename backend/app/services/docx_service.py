"""
DOCX Service
Handles reading .docx templates, extracting {key} / {key: "explain"} placeholders,
replacing them with generated content, and building output filenames.
"""
import re
import base64
import io
from typing import Dict, Tuple
from datetime import datetime

from docx import Document
from docx.oxml.ns import qn
from app.core.logging_config import get_logger

logger = get_logger(__name__)

# Regex: matches {anything}, including colons and quotes inside
PLACEHOLDER_PATTERN = re.compile(r'\{([^{}]+?)\}')


def _parse_placeholder(raw: str) -> Tuple[str, str]:
    """
    Parse a raw placeholder string (inside braces) into (key, instruction).
    Examples:
      "skills"                        -> ("skills", "")
      "projects: list 3 key projects" -> ("projects", "list 3 key projects")
      'Content of Cover letter: 4 paragraphs, 1 page' -> ("Content of Cover letter", "4 paragraphs, 1 page")
    """
    # Split only on the FIRST colon
    parts = raw.split(':', 1)
    key = parts[0].strip().strip('"').strip("'")
    instruction = parts[1].strip().strip('"').strip("'") if len(parts) > 1 else ""
    return key, instruction


def extract_placeholders(docx_bytes: bytes) -> Dict[str, str]:
    """
    Parse a .docx binary and find all {key} or {key: explain} tokens.
    Searches paragraphs and all table cells.
    Returns a dict: { key: instruction }  (instruction may be empty string)
    """
    doc = Document(io.BytesIO(docx_bytes))
    found: Dict[str, str] = {}

    def scan_text(text: str):
        for match in PLACEHOLDER_PATTERN.finditer(text):
            raw = match.group(1)
            key, instruction = _parse_placeholder(raw)
            if key and key not in found:
                found[key] = instruction

    # Scan all paragraphs
    for para in doc.paragraphs:
        scan_text(para.text)

    # Scan all table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    scan_text(para.text)

    logger.info(f"Extracted {len(found)} placeholders from docx: {list(found.keys())}")
    return found


def _replace_in_paragraph(para, replacements: Dict[str, str]):
    """
    Replace all {key} / {key: explain} tokens in a paragraph in-place.
    Operates on the paragraph's full text, then rewrites the first run
    and clears the rest (to preserve formatting of the first run).
    """
    full_text = para.text
    if not PLACEHOLDER_PATTERN.search(full_text):
        return  # Nothing to replace in this paragraph

    def replace_match(m):
        raw = m.group(1)
        key, _ = _parse_placeholder(raw)
        return replacements.get(key, m.group(0))  # Keep original if no replacement

    new_text = PLACEHOLDER_PATTERN.sub(replace_match, full_text)
    if new_text == full_text:
        return  # No actual change

    # Preserve the formatting of the first run, rewrite its text, clear others
    if para.runs:
        # Copy format from first run, set its text to new_text
        first_run = para.runs[0]
        first_run.text = new_text
        # Clear remaining runs
        for run in para.runs[1:]:
            run.text = ""
    else:
        # No runs — rare, but handle it
        para.text = new_text


def replace_placeholders(docx_bytes: bytes, replacements: Dict[str, str]) -> bytes:
    """
    Replace all {key} / {key: explain} tokens in the .docx with the replacement values.
    Preserves document structure and formatting.
    Returns the modified .docx as bytes.
    """
    doc = Document(io.BytesIO(docx_bytes))

    # Replace in body paragraphs
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    # Replace in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)

    # Replace in headers and footers
    for section in doc.sections:
        for hdr_para in section.header.paragraphs:
            _replace_in_paragraph(hdr_para, replacements)
        for ftr_para in section.footer.paragraphs:
            _replace_in_paragraph(ftr_para, replacements)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


# Filename placeholder patterns (supports English and Chinese variants)
_FILENAME_REPLACEMENTS = [
    # Pattern -> replaced by
    (re.compile(r'ResumeT\b'), 'Resume'),
    (re.compile(r'Cover LetterT\b'), 'Cover Letter'),
    (re.compile(r'\{Company name\}', re.IGNORECASE), '{company}'),
    (re.compile(r'\{Position name\}', re.IGNORECASE), '{position}'),
    (re.compile(r'\{Name of applicant\}', re.IGNORECASE), '{applicant}'),
    (re.compile(r'\{Generate date\}', re.IGNORECASE), '{date}'),
    # Chinese equivalents
    (re.compile(r'\{公司名称\}'), '{company}'),
    (re.compile(r'\{职位名称\}'), '{position}'),
    (re.compile(r'\{姓名\}'), '{applicant}'),
]


def build_output_filename(
    company: str,
    position: str,
    applicant: str = "",
    date_str: str = "",
    doc_type: str = "Resume"
) -> str:
    """
    Constructs a standardized output filename.
    Format: "{doc_type} - {company} - {position} - {applicant} - {date_str}.docx"
    """
    if not date_str:
        date_str = datetime.now().strftime("%Y%m%d")

    # Sanitize values for use in filename
    def safe(s: str) -> str:
        s = s or ""
        return re.sub(r'[<>:"/\\|?*]', '', s).strip()

    safe_company = safe(company)
    safe_position = safe(position)
    safe_applicant = safe(applicant)

    parts = [doc_type]
    if safe_company: parts.append(safe_company)
    if safe_position: parts.append(safe_position)
    if safe_applicant: parts.append(safe_applicant)
    if date_str: parts.append(date_str)

    result = " - ".join(parts) + ".docx"

    logger.info(f"Built output filename: {result}")
    return result
